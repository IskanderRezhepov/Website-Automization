from __future__ import annotations

from dataclasses import dataclass, field as dataclass_field
from pathlib import Path
from typing import Literal
from concurrent.futures import ThreadPoolExecutor, as_completed

import hashlib
import json
import os
import re
from functools import lru_cache

import fitz
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pytesseract import Output

from .text_utils import normalize_text

SUPPORTED_EXTENSIONS = {'.pdf', '.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp', '.webp', '.docx', '.txt'}


@dataclass
class PageContent:
    page_number: int
    text: str
    extraction_method: Literal['digital', 'ocr', 'hybrid', 'none']
    char_count: int
    quality: float
    variants: dict[str, str] = dataclass_field(default_factory=dict)
    cache_hit: bool = False
    layout_words: list[dict] = dataclass_field(default_factory=list)
    page_width: float | None = None
    page_height: float | None = None
    analysis_profile: str | None = None


@dataclass
class ReadDocument:
    filename: str
    page_count: int
    source_type: str
    pages: list[PageContent]

    @property
    def full_text(self) -> str:
        return '\n'.join(page.text for page in self.pages)

    @property
    def used_ocr(self) -> bool:
        return any(page.extraction_method in {'ocr', 'hybrid'} for page in self.pages)


class OCRUnavailableError(RuntimeError):
    pass


@lru_cache(maxsize=8)
def _effective_ocr_languages(requested: str) -> str:
    """Use installed requested languages and retain an English fallback."""
    try:
        available = set(pytesseract.get_languages(config=""))
    except pytesseract.TesseractNotFoundError as exc:
        raise OCRUnavailableError(
            "Tesseract OCR не установлен или не добавлен в PATH."
        ) from exc
    selected = [
        language for language in str(requested or "").split("+")
        if language in available
    ]
    if selected:
        return "+".join(selected)
    if "eng" in available:
        return "eng"
    if available:
        return sorted(available)[0]
    raise OCRUnavailableError("В Tesseract не установлены языковые модели OCR.")


FILE_CACHE_VERSION = "document-v3"


def _document_cache_path(path: Path, languages: str, dpi: int, minimum: int,
                         cache_dir: Path | None, mode: str) -> Path | None:
    if cache_dir is None:
        return None
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    digest.update(f"|{languages}|{dpi}|{minimum}|{mode}|{FILE_CACHE_VERSION}".encode())
    return cache_dir / "documents" / f"{digest.hexdigest()}.json"


def _load_document_cache(cache_path: Path | None, filename: str) -> ReadDocument | None:
    if cache_path is None or not cache_path.exists():
        return None
    try:
        payload = json.loads(cache_path.read_text(encoding="utf-8"))
        pages = []
        for raw in payload["pages"]:
            raw = dict(raw)
            if raw.get("extraction_method") in {"ocr", "hybrid"}:
                raw["cache_hit"] = True
                raw["analysis_profile"] = f"{raw.get('analysis_profile') or 'ocr'}-cached"
            pages.append(PageContent(**raw))
        return ReadDocument(filename, int(payload["page_count"]),
                            str(payload.get("source_type") or "pdf"), pages)
    except (OSError, ValueError, TypeError, KeyError):
        return None


def _save_document_cache(cache_path: Path | None, document: ReadDocument) -> None:
    if cache_path is None:
        return
    payload = {
        "page_count": document.page_count,
        "source_type": document.source_type,
        "pages": [
            {
                "page_number": page.page_number,
                "text": page.text,
                "extraction_method": page.extraction_method,
                "char_count": page.char_count,
                "quality": page.quality,
                "variants": page.variants,
                "cache_hit": page.cache_hit,
                "layout_words": page.layout_words,
                "page_width": page.page_width,
                "page_height": page.page_height,
                "analysis_profile": page.analysis_profile,
            }
            for page in document.pages
        ],
    }
    try:
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        temporary = cache_path.with_suffix(f".{os.getpid()}.tmp")
        temporary.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        temporary.replace(cache_path)
    except OSError:
        pass


def read_document(path: Path, *, ocr_languages: str, ocr_dpi: int, min_digital_chars: int, ocr_cache_dir: Path | None = None, analysis_mode: str = 'standard') -> ReadDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f'Неподдерживаемый формат: {suffix}')
    if suffix == '.pdf':
        cache_path = _document_cache_path(
            path, ocr_languages, ocr_dpi, min_digital_chars,
            ocr_cache_dir, analysis_mode,
        )
        cached = _load_document_cache(cache_path, path.name)
        if cached is not None:
            return cached
        result = _read_pdf(path, ocr_languages, ocr_dpi, min_digital_chars, ocr_cache_dir, analysis_mode)
        _save_document_cache(cache_path, result)
        return result
    if suffix == '.docx':
        return _read_docx(path)
    if suffix == '.txt':
        text = normalize_text(path.read_text(encoding='utf-8', errors='ignore'))
        return ReadDocument(path.name, 1, 'text', [PageContent(1, text, 'digital', len(text), 0.99, {'digital': text})])
    return _read_image(path, ocr_languages, analysis_mode)


def _read_docx(path: Path) -> ReadDocument:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(' | '.join(cell.text.strip() for cell in row.cells))
    text = normalize_text('\n'.join(parts))
    return ReadDocument(path.name, 1, 'docx', [PageContent(1, text, 'digital', len(text), 0.99, {'digital': text})])



def _digital_layout_words(page: fitz.Page) -> list[dict]:
    words = []
    for item in page.get_text("words") or []:
        if len(item) < 5:
            continue
        x0, y0, x1, y1, value = item[:5]
        value = str(value).strip()
        if not value:
            continue
        words.append({
            "text": value,
            "x0": round(float(x0), 2),
            "y0": round(float(y0), 2),
            "x1": round(float(x1), 2),
            "y1": round(float(y1), 2),
            "confidence": 1.0,
        })
    return words


def _choose_auto_ocr(image: Image.Image, languages: str, dpi: int,
                     cache_dir: Path | None) -> tuple[dict, str]:
    """Choose OCR effort separately for one page.

    The first pass is intentionally cheap. A stronger pass runs only when
    the page remains weak, short or structurally suspicious.
    """
    fast = _ocr_cached(image, languages, dpi, cache_dir, "fast")
    fast_text = normalize_text(fast.get("preferred", ""))
    fast_quality = float(fast.get("quality") or 0)

    layout_count = len(fast.get("layout_words") or [])
    # Mean OCR confidence is often modest on otherwise complete bilingual
    # contracts.  Low confidence alone previously forced a full second OCR of
    # every page.  Escalate when content is actually missing/structurally weak,
    # or when both confidence and useful text volume are poor.
    needs_stronger = (
        len(fast_text) < 70
        or layout_count < 8
        or (fast_quality < 0.35 and len(fast_text) < 140)
    )
    if not needs_stronger:
        return fast, "auto-fast"

    # The previous implementation escalated directly from ``fast`` to
    # ``accurate``.  Accurate mode can run five Tesseract passes, so one weak
    # scanned page could be OCRed six times.  Auto mode is the normal/default
    # workflow and should be bounded: use a two/three-pass standard fallback.
    # Users can still explicitly choose ``accurate`` for a difficult page.
    fallback = _ocr_cached(image, languages, max(dpi, 190), cache_dir, "standard")
    fallback_text = normalize_text(fallback.get("preferred", ""))
    fallback_quality = float(fallback.get("quality") or 0)

    fast_score = fast_quality + min(len(fast_text), 1200) / 5000
    fallback_score = fallback_quality + min(len(fallback_text), 1200) / 5000
    if fallback_score >= fast_score:
        return fallback, "auto-standard"
    return fast, "auto-fast"



def _read_pdf(path: Path, languages: str, dpi: int, min_digital_chars: int, cache_dir: Path | None = None, analysis_mode: str = 'standard') -> ReadDocument:
    pages_by_number: dict[int, PageContent] = {}
    ocr_jobs: list[dict] = []
    with fitz.open(path) as doc:
        if doc.page_count <= 0:
            raise RuntimeError(
                "PDF повреждён или не содержит доступных страниц. "
                "Повторно сохраните файл как PDF и загрузите исправленную копию."
            )
        for index, page in enumerate(doc, start=1):
            digital = normalize_text(page.get_text('text') or '')
            digital_words = _digital_layout_words(page)
            page_width = float(page.rect.width)
            page_height = float(page.rect.height)
            digital_quality = _digital_text_quality(digital)
            digital_is_usable = _digital_text_is_usable(digital, min_digital_chars)
            # In auto mode, short signature/specification pages with clean
            # embedded text should not trigger multiple expensive OCR passes.
            if analysis_mode == "auto" and len(digital) >= 80 and digital_quality >= 0.82:
                digital_is_usable = True
            if analysis_mode == "auto" and len(digital_words) >= 35 and digital_quality >= 0.72:
                digital_is_usable = True
            # Multi-page contracts often end with Documentolog signature
            # sheets.  Their embedded text is short and dominated by IDs /
            # Base64 signatures, which gives a low language-quality score even
            # though the text layer is complete. OCRing those pages can take
            # several minutes and adds no useful contractual evidence.
            if (
                analysis_mode == "auto"
                and doc.page_count > 1
                and pages_by_number
                and len(digital) >= 200
                and len(digital_words) >= 25
                and (
                    "ЭЛЕКТРОННЫЙ ДОКУМЕНТ ПОДПИСАН" in digital.upper()
                    or "DOCUMENTOLOG" in digital.upper()
                    or "ДАТА ПОДПИСАНИЯ" in digital.upper()
                )
            ):
                digital_is_usable = True
            # Long bilingual legal PDFs can score below the generic quality
            # threshold because of narrow columns, Latin company names and
            # many contract identifiers, even though the embedded text is
            # perfectly usable.  Prefer the digital layer when it contains
            # substantial Cyrillic prose and several legal anchors.
            if analysis_mode == "auto" and not digital_is_usable:
                upper_digital = digital.upper()
                cyrillic_chars = sum("\u0400" <= ch <= "\u052f" for ch in digital)
                legal_hits = sum(token in upper_digital for token in (
                    "ДОГОВОР", "ЛИЗИНГ", "СТОРОН", "ТЕНГЕ",
                    "БИН", "ИИН", "ШАРТ", "СОГЛАШЕНИ",
                ))
                if (
                    len(digital) >= 450
                    and len(digital_words) >= 70
                    and cyrillic_chars >= 180
                    and legal_hits >= 3
                ):
                    digital_is_usable = True
            malformed_iban_layer = bool(re.search(
                r"(?:KZ|КZ|KЗ|КЗ)\s*\d{8,17}\b|"
                r"(?:КZ|KЗ|КЗ)\s*\d{18}\b",
                digital,
                re.I,
            ))
            if digital_is_usable and not malformed_iban_layer:
                pages_by_number[index] = PageContent(
                    index, digital, 'digital', len(digital),
                    round(_digital_text_quality(digital), 2), {'digital': digital}, False,
                    digital_words, page_width, page_height, 'digital',
                )
                continue

            effective_dpi = max(dpi, 280) if malformed_iban_layer else dpi
            image = _render_page(page, effective_dpi)
            ocr_jobs.append({
                "index": index, "digital": digital, "image": image,
                "force_accurate": malformed_iban_layer,
                "effective_dpi": effective_dpi,
            })
        page_count = doc.page_count

    # Tesseract runs as an external process, so page-level threads execute
    # independently despite Python's GIL. Bound concurrency to protect smaller
    # laptops while reducing a four-page scanned contract to roughly one page's
    # wall-clock time.
    def run_job(job):
        if job.get("force_accurate"):
            ocr = _ocr_cached(
                job["image"], languages, job.get("effective_dpi") or dpi,
                cache_dir, "accurate",
            )
            page_profile = "auto-iban-accurate"
        elif analysis_mode == "auto":
            ocr, page_profile = _choose_auto_ocr(job["image"], languages, dpi, cache_dir)
        else:
            ocr = _ocr_cached(job["image"], languages, dpi, cache_dir, analysis_mode)
            page_profile = analysis_mode
        return job, ocr, page_profile

    if ocr_jobs:
        workers = min(len(ocr_jobs), max(1, min(4, (os.cpu_count() or 2))))
        with ThreadPoolExecutor(max_workers=workers, thread_name_prefix="ocr-page") as pool:
            futures = [pool.submit(run_job, job) for job in ocr_jobs]
            for future in as_completed(futures):
                job, ocr, page_profile = future.result()
                digital = job["digital"]
                image = job["image"]
                combined = normalize_text(f'{digital}\n{ocr["preferred"]}' if digital else ocr['preferred'])
                method = 'hybrid' if digital and combined else 'ocr' if combined else 'none'
                pages_by_number[job["index"]] = PageContent(
                    job["index"], combined, method, len(combined), ocr['quality'],
                    ocr['variants'], bool(ocr.get('cache_hit')),
                    ocr.get('layout_words', []), float(image.width),
                    float(image.height), page_profile,
                )

    pages = [pages_by_number[index] for index in range(1, page_count + 1)]
    return ReadDocument(path.name, page_count, 'pdf', pages)


def _read_image(path: Path, languages: str, analysis_mode: str = 'standard') -> ReadDocument:
    try:
        image = ImageOps.exif_transpose(Image.open(path)).convert('RGB')
        if analysis_mode == "auto":
            fast = _ocr_multimode(image, languages, "fast")
            if len(normalize_text(fast.get("preferred", ""))) < 180 or float(fast.get("quality") or 0) < 0.62:
                accurate = _ocr_multimode(image, languages, "accurate")
                fast_score = float(fast.get("quality") or 0) + min(len(fast.get("preferred", "")), 1200) / 5000
                accurate_score = float(accurate.get("quality") or 0) + min(len(accurate.get("preferred", "")), 1200) / 5000
                ocr = accurate if accurate_score >= fast_score else fast
                page_profile = "auto-accurate" if ocr is accurate else "auto-fast"
            else:
                ocr = fast
                page_profile = "auto-fast"
        else:
            ocr = _ocr_multimode(image, languages, analysis_mode)
            page_profile = analysis_mode
    except Exception as exc:
        raise RuntimeError(f'Не удалось прочитать изображение: {exc}') from exc
    text = normalize_text(ocr['preferred'])
    return ReadDocument(path.name, 1, 'image', [PageContent(1, text, 'ocr', len(text), ocr['quality'], ocr['variants'], bool(ocr.get('cache_hit')), ocr.get('layout_words', []), float(image.width), float(image.height), page_profile)])


def _digital_text_quality(text: str) -> float:
    """
    Estimate whether the embedded PDF text is genuinely readable.

    Some scanned PDFs contain a broken hidden OCR layer with thousands of
    characters such as ``3aJIOrO...``. Character count alone incorrectly marks
    this as perfect digital text. The score below is deliberately tailored to
    Russian/Kazakh legal documents while still allowing normal Latin names,
    contract numbers and IBANs.
    """
    if not text:
        return 0.0

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return 0.0

    cyrillic = sum(
        "\u0400" <= ch <= "\u052f"
        or ch in "ӘәҒғҚқҢңӨөҰұҮүІіҺһ"
        for ch in letters
    )
    latin = sum(("A" <= ch <= "Z") or ("a" <= ch <= "z") for ch in letters)

    tokens = re.findall(r"\b[\w-]{4,}\b", text, re.UNICODE)
    mixed_alnum = sum(
        bool(re.search(r"[A-Za-zА-Яа-яӘәҒғҚқҢңӨөҰұҮүІіҺһ]", token))
        and bool(re.search(r"\d", token))
        for token in tokens
    )
    mixed_ratio = mixed_alnum / max(len(tokens), 1)

    # Broken OCR frequently contains long Latin-looking transliterations of
    # Cyrillic words. Real banking documents still contain much more Cyrillic
    # prose than Latin prose.
    latin_heavy = latin > cyrillic * 0.75 and cyrillic < len(letters) * 0.55

    common_terms = (
        "ДОГОВОР", "СОГЛАШЕНИ", "СТОРОН", "БИН", "ИИН", "ТЕНГЕ",
        "СЧЕТ", "СЧЁТ", "ЗАЛОГ", "ЛИЗИНГ", "КЕПІЛ", "ШАРТ",
    )
    upper = text.upper()
    language_hits = sum(term in upper for term in common_terms)

    base = min(1.0, len(text) / 1200)
    alphabetic_ratio = len(letters) / max(len(text), 1)
    score = 0.45 * base + 0.35 * min(1.0, alphabetic_ratio / 0.45)
    score += min(0.20, language_hits * 0.035)

    if mixed_ratio > 0.08:
        score -= 0.35
    if latin_heavy:
        score -= 0.30

    return max(0.0, min(1.0, score))


def _digital_text_is_usable(text: str, minimum: int) -> bool:
    if len(text) < minimum:
        return False
    return _digital_text_quality(text) >= 0.62


def _render_page(page: fitz.Page, dpi: int) -> Image.Image:
    zoom = dpi / 72
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    return Image.frombytes('RGB', [pix.width, pix.height], pix.samples)


def _preprocess_variants(image: Image.Image) -> dict[str, Image.Image]:
    gray = ImageOps.exif_transpose(image).convert('L')
    gray = ImageOps.autocontrast(gray, cutoff=1)
    contrast = ImageEnhance.Contrast(gray).enhance(1.45).filter(ImageFilter.SHARPEN)
    threshold = contrast.point(lambda p: 255 if p > 175 else 0)
    return {'gray': gray, 'contrast': contrast, 'threshold': threshold}



def _ocr_cached(image: Image.Image, languages: str, dpi: int, cache_dir: Path | None, analysis_mode: str = 'standard') -> dict:
    if cache_dir is None:
        result = _ocr_multimode(image, languages, analysis_mode)
        result['cache_hit'] = False
        return result

    cache_dir.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256()
    digest.update(image.tobytes())
    digest.update(f"|{languages}|{dpi}|adaptive-v2|{analysis_mode}".encode("utf-8"))
    cache_path = cache_dir / f"{digest.hexdigest()}.json"

    if cache_path.exists():
        try:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
            if isinstance(cached, dict) and "preferred" in cached:
                cached["cache_hit"] = True
                return cached
        except (OSError, ValueError, TypeError):
            pass

    result = _ocr_multimode(image, languages, analysis_mode)
    result['cache_hit'] = False
    try:
        cache_path.write_text(json.dumps(result, ensure_ascii=False), encoding="utf-8")
    except OSError:
        pass
    return result



def _ocr_multimode(image: Image.Image, languages: str, analysis_mode: str = 'standard') -> dict:
    """
    Fast adaptive OCR:
    1. one general pass for every page;
    2. a table pass only when the first pass is weak;
    3. sparse/column passes only for difficult pages.

    This is substantially faster than running 3-5 Tesseract passes on every
    page, especially for long scanned contracts.
    """
    prepared = _preprocess_variants(image)
    width, height = image.size
    variants: dict[str, str] = {}
    qualities: list[float] = []

    mode = analysis_mode if analysis_mode in {"fast", "standard", "accurate"} else "standard"

    full_text, full_quality, full_layout = _ocr_with_quality(prepared["contrast"], languages, psm=3)
    variants["full"] = normalize_text(full_text)
    qualities.append(full_quality)

    enough_text = len(variants["full"]) >= (240 if mode == "fast" else 350)
    strong_page = full_quality >= (0.61 if mode == "fast" else 0.69) and enough_text

    # Fast mode intentionally performs one pass unless the page is almost empty.
    run_table = (
        mode == "accurate"
        or (mode == "standard" and not strong_page)
        or (mode == "fast" and len(variants["full"]) < 90)
    )
    if run_table:
        table_text, table_quality, _table_layout = _ocr_with_quality(prepared["threshold"], languages, psm=6)
        variants["table"] = normalize_text(table_text)
        qualities.append(table_quality)

    best_quality = max(qualities or [0.0])
    best_chars = max((len(value) for value in variants.values()), default=0)

    run_sparse = (
        mode == "accurate"
        or (mode == "standard" and (best_quality < 0.60 or best_chars < 220))
        or (mode == "fast" and best_chars < 70)
    )
    if run_sparse:
        sparse_text, sparse_quality, _sparse_layout = _ocr_with_quality(prepared["gray"], languages, psm=11)
        variants["sparse"] = normalize_text(sparse_text)
        qualities.append(sparse_quality)

    # Column OCR is the most expensive pass. It is disabled in fast mode,
    # conditional in standard mode and more permissive in accurate mode.
    column_threshold = 0.74 if mode == "accurate" else 0.66
    # Cropping and OCRing both columns adds two full Tesseract processes.
    # Reserve that expensive strategy for the explicitly requested accurate
    # mode; standard/auto can retry an individual page from the review screen.
    if mode == "accurate" and width / max(height, 1) > 0.78 and max(qualities or [0.0]) < column_threshold:
        for side, box in {
            "right_column": (int(width * 0.47), 0, width, height),
            "left_column": (0, 0, int(width * 0.53), height),
        }.items():
            crop = prepared["contrast"].crop(box)
            side_text, side_quality, _side_layout = _ocr_with_quality(crop, languages, psm=6)
            variants[side] = normalize_text(side_text)
            qualities.append(side_quality)

    preferred_parts = [
        variants.get("full", ""),
        variants.get("table", ""),
        variants.get("right_column", ""),
        variants.get("left_column", ""),
        variants.get("sparse", ""),
    ]
    preferred = normalize_text("\n".join(part for part in preferred_parts if part))
    return {
        "preferred": preferred,
        "quality": round(max(qualities or [0.0]), 2),
        "variants": variants,
        "layout_words": full_layout,
    }

def _ocr_with_quality(image: Image.Image, languages: str, psm: int) -> tuple[str, float, list[dict]]:
    try:
        languages = _effective_ocr_languages(languages)
        config = f'--oem 3 --psm {psm} -c preserve_interword_spaces=1'
        data = pytesseract.image_to_data(image, lang=languages, config=config, output_type=Output.DICT)
        words: list[str] = []
        layout_words: list[dict] = []
        confidences: list[float] = []
        count = len(data.get('text', []))
        for index in range(count):
            value = str(data.get('text', [''])[index]).strip()
            try:
                conf_value = float(data.get('conf', [-1])[index])
            except (TypeError, ValueError, IndexError):
                conf_value = -1
            if not value:
                continue
            words.append(value)
            if conf_value >= 0:
                confidences.append(conf_value)
            left = float(data.get('left', [0])[index])
            top = float(data.get('top', [0])[index])
            width = float(data.get('width', [0])[index])
            height = float(data.get('height', [0])[index])
            layout_words.append({
                'text': value, 'x0': left, 'y0': top,
                'x1': left + width, 'y1': top + height,
                'confidence': round(max(conf_value, 0) / 100, 3),
            })
        quality = (sum(confidences) / len(confidences) / 100) if confidences else 0.0
        return ' '.join(words), quality, layout_words
    except pytesseract.TesseractNotFoundError as exc:
        raise OCRUnavailableError('Tesseract OCR не установлен или не добавлен в PATH.') from exc
